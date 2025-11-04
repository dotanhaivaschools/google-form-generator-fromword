import re
import json
from docx import Document
from googleapiclient.discovery import build
from google.oauth2 import service_account
import streamlit as st


def get_credentials():
    """Lấy credentials từ Streamlit Secrets (GOOGLE_CREDENTIALS)"""
    info = json.loads(st.secrets["GOOGLE_CREDENTIALS"])
    return service_account.Credentials.from_service_account_info(info)


def parse_docx(file_path):
    """
    Đọc file Word (.docx) và trích xuất danh sách câu hỏi.
    Hỗ trợ:
    - File Word xuất từ Markdown (nội dung 1 đoạn)
    - Tự động tách Câu hỏi / Đáp án
    - Nhận diện A., B., C., D. dù dính liền hoặc có khoảng trắng
    - Nhận diện đáp án đúng từ ký tự được gạch chân (underline)
    - Duyệt cả đoạn văn và bảng (table)
    """
    doc = Document(file_path)
    questions = []
    current_question = None

    def extract_from_text_block(block_text, para=None):
        """Phân tích 1 đoạn văn bản, tách câu hỏi và đáp án"""
        nonlocal current_question, questions

        # Chuẩn hóa văn bản: loại bỏ xuống dòng mềm, tab, khoảng trắng thừa
        block_text = block_text.replace("\r", " ").replace("\n", " ").replace("\t", " ")
        block_text = re.sub(r"\s{2,}", " ", block_text).strip()

        # Có thể chứa nhiều câu hỏi trong 1 đoạn
        segments = re.split(r"(?=Câu\s*\d+\s*[:\.])", block_text)
        for seg in segments:
            seg = seg.strip()
            if not seg:
                continue

            # 🟢 Bắt đầu câu hỏi mới
            if re.match(r"^Câu\s*\d+\s*[:\.]", seg):
                if current_question and current_question.get("options"):
                    questions.append(current_question)
                current_question = {
                    "question": "",
                    "options": [],
                    "answer_key": ""
                }

                # Lấy nội dung câu hỏi
                match_q = re.match(r"^Câu\s*\d+\s*[:\.]\s*(.+)", seg)
                if match_q:
                    current_question["question"] = match_q.group(1).strip()

            # 🟠 Tách các đáp án A., B., C., D.
            parts = re.split(r"(?=\b[A-D]\s*\.)", seg)
            for part in parts:
                part = part.strip()
                if not re.match(r"^[A-D]\s*\.", part):
                    continue

                label_match = re.match(r"^([A-D])\s*\.", part)
                if not label_match:
                    continue
                label = label_match.group(1)
                raw_option = re.sub(r"^[A-D]\s*\.\s*", "", part).strip()

                if not raw_option:
                    raw_option = f"Tùy chọn {len(current_question['options']) + 1}"

                # 🔵 Kiểm tra gạch chân trong run (nếu đoạn văn có tham chiếu)
                if para:
                    for run in para.runs:
                        if run.underline and f"{label}." in run.text:
                            current_question["answer_key"] = raw_option

                current_question["options"].append(raw_option)

    def handle_paragraphs(paragraphs):
        """Duyệt qua danh sách paragraph và xử lý từng đoạn"""
        for para in paragraphs:
            text = para.text.strip()
            if text:
                extract_from_text_block(text, para)

    # ✅ Duyệt toàn bộ đoạn văn chính
    handle_paragraphs(doc.paragraphs)

    # ✅ Duyệt cả bảng (table) nếu có
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                handle_paragraphs(cell.paragraphs)

    # ✅ Thêm câu cuối cùng (nếu có)
    if current_question and current_question.get("options"):
        questions.append(current_question)

    # ⚠️ Nếu không có câu hỏi nào
    if not questions:
        raise ValueError(
            "Không trích xuất được câu hỏi nào từ file Word! "
            "Hãy đảm bảo rằng mỗi câu hỏi bắt đầu bằng 'Câu n:' "
            "và có ít nhất một đáp án A."
        )

    return questions


def create_google_form(questions, form_title, share_email=None):
    """Tạo Google Form từ danh sách câu hỏi"""
    credentials = get_credentials()
    service = build("forms", "v1", credentials=credentials)

    # 🧾 Tạo biểu mẫu Google Form
    form = {
        "info": {
            "title": form_title,
            "documentTitle": form_title
        }
    }
    result = service.forms().create(body=form).execute()
    form_id = result["formId"]

    requests = []

    for q in questions:
        cleaned = [opt.strip().replace("\n", " ") for opt in q["options"] if opt.strip()]
        unique_options = list(dict.fromkeys(cleaned))

        if not unique_options:
            continue

        # ⭐ Gắn sao cho đáp án đúng
        labeled_options = []
        for opt in unique_options:
            if opt == q.get("answer_key", "").replace("\n", " ").strip():
                labeled_options.append(f"{opt} ⭐")
            else:
                labeled_options.append(opt)

        question_title = q["question"].replace("\n", " ").strip()

        question_item = {
            "createItem": {
                "item": {
                    "title": question_title,
                    "questionItem": {
                        "question": {
                            "required": True,
                            "choiceQuestion": {
                                "type": "RADIO",
                                "options": [{"value": opt} for opt in labeled_options],
                                "shuffle": False
                            }
                        }
                    }
                },
                "location": {"index": 0}
            }
        }
        requests.append(question_item)

    # 📤 Gửi toàn bộ câu hỏi lên Google Form
    service.forms().batchUpdate(formId=form_id, body={"requests": requests}).execute()

    # 📬 Cấp quyền chỉnh sửa nếu có email
    if share_email:
        drive_service = build("drive", "v3", credentials=credentials)
        drive_service.permissions().create(
            fileId=form_id,
            body={
                "type": "user",
                "role": "writer",
                "emailAddress": share_email
            },
            sendNotificationEmail=True
        ).execute()

    return f"https://docs.google.com/forms/d/{form_id}/edit"